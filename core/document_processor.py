"""
文档处理模块
负责 PDF/Word 文档的解析和文本切分
"""

import os
import re
import tempfile
from pathlib import Path
from typing import List, Optional, BinaryIO

from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter

import sys
sys.path.append(str(Path(__file__).parent.parent))
from config import CHUNK_SIZE, CHUNK_OVERLAP, UPLOADS_DIR


class DocumentProcessor:
    """文档处理器：解析和切分文档"""
    
    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP
    ):
        """
        初始化文档处理器
        
        Args:
            chunk_size: 每个文本块的最大字符数
            chunk_overlap: 文本块之间的重叠字符数
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # 初始化文本切分器
        # 使用多种分隔符递归切分，优先保持段落和句子完整性
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                "\n\n",  # 段落分隔
                "\n",    # 换行
                "。",    # 中文句号
                "！",    # 中文感叹号
                "？",    # 中文问号
                ".",     # 英文句号
                "!",     # 英文感叹号
                "?",     # 英文问号
                ";",     # 分号
                "；",    # 中文分号
                " ",     # 空格
                ""       # 字符级别（最后手段）
            ]
        )
    
    def load_pdf(self, file_path: str) -> List[Document]:
        """
        加载 PDF 文件
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            Document 列表，每页一个 Document
        """
        loader = PyPDFLoader(file_path)
        documents = loader.load()
        
        # 为每个文档添加文件名元数据
        file_name = Path(file_path).name
        for doc in documents:
            doc.metadata["source_file"] = file_name
            # 页码从 1 开始（原始是从 0 开始）
            if "page" in doc.metadata:
                doc.metadata["page"] = doc.metadata["page"] + 1
        
        return documents
    
    def load_pdf_from_upload(self, uploaded_file: BinaryIO, filename: str) -> List[Document]:
        """
        从上传的文件对象加载 PDF
        
        Args:
            uploaded_file: Streamlit 上传的文件对象
            filename: 原始文件名
            
        Returns:
            Document 列表
        """
        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            documents = self.load_pdf(tmp_path)
            # 更新 source_file 为原始文件名
            for doc in documents:
                doc.metadata["source_file"] = filename
            return documents
        finally:
            # 删除临时文件
            os.unlink(tmp_path)
    
    def load_word(self, file_path: str) -> List[Document]:
        """
        加载 Word 文档 (.docx)
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            Document 列表
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
        
        docx_doc = DocxDocument(file_path)
        file_name = Path(file_path).name
        
        # 提取所有段落文本
        full_text = []
        for para in docx_doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 创建一个 Document 对象
        content = "\n\n".join(full_text)
        document = Document(
            page_content=content,
            metadata={
                "source_file": file_name,
                "page": 1  # Word 文档不分页，统一设为 1
            }
        )
        
        return [document]
    
    def load_word_from_upload(self, uploaded_file: BinaryIO, filename: str) -> List[Document]:
        """
        从上传的文件对象加载 Word 文档
        
        Args:
            uploaded_file: Streamlit 上传的文件对象
            filename: 原始文件名
            
        Returns:
            Document 列表
        """
        suffix = ".docx" if filename.endswith(".docx") else ".doc"
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            documents = self.load_word(tmp_path)
            for doc in documents:
                doc.metadata["source_file"] = filename
            return documents
        finally:
            os.unlink(tmp_path)
    
    def clean_text(self, text: str) -> str:
        """
        清洗文本：移除参考文献、致谢、页眉页脚等噪音
        同时处理 PDF 解析可能产生的无效 Unicode 字符
        """
        # ★ 移除无效的 Unicode 代理字符（surrogates）
        try:
            text = text.encode('utf-8', 'surrogatepass').decode('utf-8', 'ignore')
        except Exception:
            text = text.encode('utf-8', 'ignore').decode('utf-8', 'ignore')
        
        # ★★ 移除 PDF 解析产生的 Unicode 转义序列垃圾
        # 匹配 /uniXXXXXXXX 或 /uni0000XXXX 等模式
        text = re.sub(r'/uni[0-9a-fA-F]{8}', '', text)
        text = re.sub(r'/uni[0-9a-fA-F]{4,}', '', text)
        
        # 移除其他常见的 PDF 解析垃圾字符序列
        text = re.sub(r'\x00+', '', text)  # NULL 字符
        text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f]', '', text)  # 控制字符
        
        # 移除连续的数字垃圾（如 00057 开头的序列）
        text = re.sub(r'\b\d{5,}\b(?:/uni[0-9a-fA-F]+)*', '', text)
        
        # 移除常见的页眉页脚模式
        text = re.sub(r'第\s*\d+\s*页', '', text)
        text = re.sub(r'(?i)page\s*\d+\s*(of\s*\d+)?', '', text)
        
        # 移除多余的空白行和空格
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {3,}', ' ', text)  # 多个连续空格
        
        return text.strip()
    
    def is_garbage_chunk(self, text: str) -> bool:
        """
        判断一个文本块是否主要是垃圾内容
        使用通用化规则，适用于各种论文格式
        
        Returns:
            True 如果是垃圾块应该过滤掉
        """
        # === 检查 0：最小有效长度 ===
        if len(text) < 100:
            preview = text[:60].replace('\n', ' ')
            print(f"[DocumentProcessor] 🗑️ 过滤过短片段 (长度={len(text)}): {preview}...")
            return True
        
        text_lower = text.lower()
        
        # === 检查 1：有意义字符比例 ===
        meaningful_chars = re.findall(r'[a-zA-Z\u4e00-\u9fff0-9\.,;:!?\'"()\[\]\-\s]', text)
        meaningful_ratio = len(meaningful_chars) / len(text) if text else 0
        
        if meaningful_ratio < 0.7:
            preview = text[:60].replace('\n', ' ')
            print(f"[DocumentProcessor] 🗑️ 过滤垃圾块 (有意义字符={meaningful_ratio:.1%}): {preview}...")
            return True
        
        # === 检查 1.5：大规模作者列表（如 GPT-4 论文有数百位作者）===
        # 特征：大量逗号分隔的人名，几乎没有动词或句子结构
        # 模式："Name Name, Name Name, Name Name,..." 或 "Name, Name, Name,..."
        comma_count = text.count(',')
        if comma_count >= 15:
            # 按逗号分割，计算平均段长
            segments = [s.strip() for s in text.split(',') if s.strip()]
            if segments:
                avg_segment_len = sum(len(s) for s in segments) / len(segments)
                # 人名特征：平均段长很短（通常 < 25 字符）
                if avg_segment_len < 25:
                    # 进一步验证：检测大写字母开头的单词密度（人名特征）
                    capitalized_words = re.findall(r'\b[A-Z][a-z]+\b', text)
                    words = text.split()
                    cap_ratio = len(capitalized_words) / len(words) if words else 0
                    
                    # 如果 > 60% 的单词以大写开头（人名特征）
                    if cap_ratio > 0.6:
                        preview = text[:60].replace('\n', ' ')
                        print(f"[DocumentProcessor] 🗑️ 过滤大规模作者列表 (逗号={comma_count}, 平均段长={avg_segment_len:.1f}, 大写比={cap_ratio:.1%}): {preview}...")
                        return True
        
        # === 检查 2：参考文献列表（通用化检测）===
        # 特征：作者姓名缩写(X. Name)、卷号页码(36(4):)、年份、期刊会议名
        ref_indicators = 0
        
        # 2a. 作者姓名缩写模式：A. Name, B. Name 或 Name, A., Name, B.
        author_initials = re.findall(r'\b[A-Z]\.\s*[A-Z][a-z]+', text)  # A. Smith
        author_initials2 = re.findall(r'[A-Z][a-z]+,\s*[A-Z]\.', text)  # Smith, A.
        if len(author_initials) + len(author_initials2) >= 3:
            ref_indicators += 2
        
        # 2b. 卷号/期号/页码模式：36(4):, vol. 12, pp. 123-456
        volume_patterns = re.findall(r'\d+\(\d+\):', text)  # 36(4):
        page_patterns = re.findall(r'(?:pp?\.|pages?)\s*\d+[-–]\d+', text_lower)  # pp. 123-456
        if len(volume_patterns) >= 1 or len(page_patterns) >= 2:
            ref_indicators += 2
        
        # 2c. 多个独立年份（如 ", 2020.", ", 2023."）
        year_with_punct = re.findall(r'[,\.]\s*(19|20)\d{2}[,\.\)]', text)
        if len(year_with_punct) >= 3:
            ref_indicators += 1
        
        # 2d. 期刊/会议关键词
        journal_keywords = [
            r'transactions\s+on', r'journal\s+of', r'proceedings\s+of',
            r'conference\s+on', r'symposium\s+on', r'workshop\s+on',
            r'in\s+proc\.', r'arxiv\s*:', r'\bin\s+the\s+\d+'
        ]
        journal_matches = sum(1 for p in journal_keywords if re.search(p, text_lower))
        if journal_matches >= 1:
            ref_indicators += 1
        
        # 2e. 连续多个逗号分隔的人名（参考文献列表特征）
        comma_names = re.findall(r'[A-Z][a-z]+,\s*[A-Z]\.,?\s*(?:and\s+)?[A-Z][a-z]+', text)
        if len(comma_names) >= 2:
            ref_indicators += 1
        
        if ref_indicators >= 3:
            preview = text[:60].replace('\n', ' ')
            print(f"[DocumentProcessor] 🗑️ 过滤参考文献列表 (指标={ref_indicators}): {preview}...")
            return True
        
        # === 检查 3：论文标题页/作者信息（通用化检测）===
        title_page_indicators = 0
        
        # 3a. 作者名+机构上标：Name1, Name2,3, Name† (数字紧跟名字)
        author_superscripts = re.findall(r'[A-Z][a-z]+\d+[,\d]*\s', text)
        if len(author_superscripts) >= 2:
            title_page_indicators += 2
        
        # 3b. 机构名+上标数字：1University, 2Google, 3Microsoft
        institution_patterns = re.findall(r'\d+[A-Z][a-z]+\s+(University|Institute|Lab|Google|Microsoft|Meta|Research)', text)
        if len(institution_patterns) >= 1:
            title_page_indicators += 2
        
        # 3c. Abstract 开头模式
        if re.search(r'\babstract\s+(we|how|this|in)\s+', text_lower):
            title_page_indicators += 1
        
        # 3d. 论文标题格式：带冒号的标题 + 机构
        paper_title_pattern = re.search(r'^[A-Z][^\.]+:\s*[A-Z][^\.]+\s+[A-Z][a-z]+\d', text)
        if paper_title_pattern:
            title_page_indicators += 2
        
        # 3e. 邮箱
        email_count = len(re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', text))
        if email_count >= 1:
            title_page_indicators += 1
        
        if title_page_indicators >= 2:
            preview = text[:60].replace('\n', ' ')
            print(f"[DocumentProcessor] 🗑️ 过滤标题页/作者信息 (指标={title_page_indicators}): {preview}...")
            return True
        
        # === 检查 4：图表标题/表格数据（通用化检测）===
        figure_table_indicators = 0
        
        # 4a. Figure/Table 标题（任意位置）
        fig_table_matches = re.findall(r'(figure|fig\.?|table)\s*\d+\s*[:\.]?', text_lower)
        if len(fig_table_matches) >= 1:
            figure_table_indicators += 1
        # 多个图表引用
        if len(fig_table_matches) >= 2:
            figure_table_indicators += 1
        
        # 4b. 子图标签 (a), (b), (c) 连续出现
        subfig_labels = re.findall(r'\([a-d]\)', text_lower)
        if len(subfig_labels) >= 3:
            figure_table_indicators += 1
        
        # 4c. 连续的小数数据（表格特征）：0.762 0.833 0.864 或 92.79 99.89 98.43
        consecutive_decimals = re.findall(r'\d+\.\d+\s+\d+\.\d+\s+\d+\.\d+', text)
        if len(consecutive_decimals) >= 1:
            figure_table_indicators += 2
        
        # 4d. 空格分隔的多个数字（表格数据）
        space_separated_nums = re.findall(r'\d+\.\d+\s+\d+\.\d+', text)
        if len(space_separated_nums) >= 3:
            figure_table_indicators += 1
        
        # 4e. ± 符号（误差范围）
        if '±' in text and re.search(r'±\s*\d+\.?\d*', text):
            figure_table_indicators += 1
        
        # 如果有图表标题和数据特征的组合
        if figure_table_indicators >= 2:
            preview = text[:60].replace('\n', ' ')
            print(f"[DocumentProcessor] 🗑️ 过滤图表/表格内容 (指标={figure_table_indicators}): {preview}...")
            return True
        
        # === 检查 5：致谢/资助信息 ===
        ack_indicators = 0
        
        # 5a. 关键词匹配
        ack_keywords = [
            r'grant\s*no\.?', r'national\s+science\s+foundation', r'supported\s+by',
            r'funded\s+by', r'this\s+work\s+was\s+supported', r'acknowledgment',
            r'acknowledgement', r'we\s+thank', r'faculty\s+award', r'research\s+award',
            r'nsf\s+', r'onr\s+', r'darpa\s+'
        ]
        ack_matches = sum(1 for p in ack_keywords if re.search(p, text_lower))
        ack_indicators += ack_matches
        
        # 5b. 资助号模式（如 N00014-22-1-2773, IIS-1234567）
        grant_numbers = re.findall(r'[A-Z]{2,}\s*[-\s]?\d{4,}[-\d]*', text)
        if len(grant_numbers) >= 1:
            ack_indicators += 2
        
        if ack_indicators >= 2:
            preview = text[:60].replace('\n', ' ')
            print(f"[DocumentProcessor] 🗑️ 过滤致谢/资助信息 (指标={ack_indicators}): {preview}...")
            return True
        
        # === 检查 6：URDF/配置文件内容 ===
        urdf_patterns = [
            r'joint_name:', r'joint_type:', r'parent_link:', r'child_link:',
            r'link_\d+', r'joint_\d+', r'<link>', r'<joint>', r'<robot>'
        ]
        urdf_matches = sum(1 for p in urdf_patterns if re.search(p, text_lower))
        if urdf_matches >= 3:
            preview = text[:60].replace('\n', ' ')
            print(f"[DocumentProcessor] 🗑️ 过滤配置文件内容 (URDF匹配={urdf_matches}): {preview}...")
            return True
        
        # === 检查 7：高密度冒号（key: value 格式）===
        colon_count = text.count(':')
        colon_density = colon_count / len(text) * 100 if text else 0
        if colon_density > 3 and colon_count > 10:
            preview = text[:60].replace('\n', ' ')
            print(f"[DocumentProcessor] 🗑️ 过滤配置内容 (冒号密度={colon_density:.1f}%): {preview}...")
            return True
        
        # === 检查 8：重复模式 ===
        repeated_pattern = re.findall(r'(\b\w+_\d+\b)', text)
        if len(repeated_pattern) > 10:
            unique_ratio = len(set(repeated_pattern)) / len(repeated_pattern)
            if unique_ratio < 0.3:
                preview = text[:60].replace('\n', ' ')
                print(f"[DocumentProcessor] 🗑️ 过滤重复模式 (唯一率={unique_ratio:.1%}): {preview}...")
                return True
        
        return False

    def remove_references_section(self, documents: List[Document]) -> List[Document]:
        """
        【强力截断】宏观过滤：
        扫描所有页面，一旦检测到 'References'、'Bibliography' 或 '参考文献' 独立标题，
        直接丢弃该页及其之后的所有页面。
        
        策略：宁可错杀不可放过，因为参考文献通常在文末。
        """
        cutoff_index = -1
        
        # 参考文献标题的正则模式列表（行首匹配）
        # 支持多种格式：纯标题、带编号（6. References, VI. References）、中文
        reference_patterns = [
            # 纯标题（精确匹配）
            r'^references\s*$',
            r'^bibliography\s*$',
            r'^reference\s*$',
            r'^参考文献\s*$',
            r'^引用文献\s*$',
            # 带阿拉伯数字编号（如 "6. References", "7 References"）
            r'^\d+\.?\s+references\s*$',
            r'^\d+\.?\s+bibliography\s*$',
            r'^\d+\.?\s*参考文献\s*$',
            # 带罗马数字编号（如 "VI. References", "V References"）
            r'^[ivxIVX]+\.?\s+references\s*$',
            r'^[ivxIVX]+\.?\s+bibliography\s*$',
            # 带方括号编号（如 "[6] References"）
            r'^\[\d+\]\s*references\s*$',
        ]
        
        # 从头开始扫描所有页面（不限制范围，确保不遗漏）
        for i, doc in enumerate(documents):
            content = doc.page_content
            lines = content.split('\n')
            
            # 检查页面的前 10 行（标题通常在页面顶部）
            for line in lines[:10]:
                clean_line = line.strip().lower()
                
                # 跳过空行
                if not clean_line:
                    continue
                
                # 尝试匹配所有模式
                for pattern in reference_patterns:
                    if re.match(pattern, clean_line, re.IGNORECASE):
                        cutoff_index = i
                        page_num = doc.metadata.get('page', i + 1)
                        print(f"[DocumentProcessor] ⚠️ 检测到参考文献标题于第 {page_num} 页: '{line.strip()}'")
                        print(f"[DocumentProcessor] 🔪 强力截断：丢弃第 {page_num} 页及之后共 {len(documents) - i} 页")
                        break
                
                if cutoff_index != -1:
                    break
            
            if cutoff_index != -1:
                break
        
        # 如果找到了参考文献页，截断该页及其之后的所有页面
        if cutoff_index != -1:
            return documents[:cutoff_index]
        
        return documents
    
    def is_reference_chunk(self, text: str) -> bool:
        """
        【超强过滤版】微观过滤：
        判断一个文本块是否主要是参考文献内容。
        使用多种特征综合判断，宁可错杀不可放过。
        """
        text_length = len(text)
        if text_length < 50:
            return False  # 太短的不处理
        
        # === 特征 1：引用标记模式 [1], [23], [1,2,3] ===
        citation_pattern = r'\[\d+(?:,\s*\d+)*\]'
        citations = re.findall(citation_pattern, text)
        
        # === 特征 2：连续编号引用检测（如 [1] xxx [2] xxx [3] xxx）===
        sequential_citations = re.findall(r'\[(\d+)\]', text)
        has_sequential = False
        if len(sequential_citations) >= 3:
            nums = [int(n) for n in sequential_citations[:10]]
            for i in range(len(nums) - 2):
                if nums[i+1] == nums[i] + 1 and nums[i+2] == nums[i] + 2:
                    has_sequential = True
                    break
        
        # === 特征 3：作者+年份格式（如 "Smith et al., 2022" 或 "Name (2021)"）===
        # 这是参考文献列表的强特征
        author_year_pattern = r'(?:[A-Z][a-z]+\s+(?:et\s+al\.?|and\s+[A-Z][a-z]+)?\s*[\(\,]\s*(?:19|20)\d{2}\s*[\)\,])'
        author_year_matches = re.findall(author_year_pattern, text)
        
        # === 特征 4：arXiv 引用模式 ===
        arxiv_refs = re.findall(r'arXiv', text, re.IGNORECASE)
        
        # === 特征 5：年份模式 (如 2023, 2021) ===
        year_pattern = r'\b(19|20)\d{2}\b'
        years = re.findall(year_pattern, text)
        
        # === 特征 6：会议/期刊关键词 ===
        venue_keywords = [
            'IEEE', 'ACM', 'CVPR', 'ICCV', 'ECCV', 'NeurIPS', 'ICML', 'ICLR',
            'AAAI', 'IJCAI', 'preprint', 'Proceedings', 'Conference', 'Journal',
            'Transactions', 'vol.', 'pp.', 'eds.', 'et al.', 'In Proceedings',
            'Workshop', 'Symposium', 'Annual Meeting', 'arXiv', 'abs/'
        ]
        venue_count = sum(1 for kw in venue_keywords if kw.lower() in text.lower())
        
        # === 特征 7：页码模式（如 pp. 123-456, pages 770–778）===
        page_patterns = re.findall(r'(?:pp\.|pages?)\s*\d+[\-–]\d+', text, re.IGNORECASE)
        
        # === 特征 8：DOI/URL 密度 ===
        doi_count = len(re.findall(r'doi[:\.\s]', text, re.IGNORECASE))
        url_count = len(re.findall(r'https?://', text))
        
        # === 综合判断逻辑（更激进的阈值） ===
        is_ref = False
        reason = ""
        
        # 规则 1：检测到连续编号引用（强特征）
        if has_sequential:
            is_ref = True
            reason = f"连续编号引用 {sequential_citations[:5]}"
        
        # 规则 2：作者+年份格式 >= 2 个
        elif len(author_year_matches) >= 2:
            is_ref = True
            reason = f"作者年份格式={len(author_year_matches)}"
        
        # 规则 3：引用标记 >= 2 个（降低阈值）
        elif len(citations) >= 2:
            is_ref = True
            reason = f"引用标记数={len(citations)}"
        
        # 规则 4：arXiv 引用 >= 1 个（降低阈值）
        elif len(arxiv_refs) >= 1:
            is_ref = True
            reason = f"arXiv引用数={len(arxiv_refs)}"
        
        # 规则 5：年份 + 会议组合（降低阈值）
        elif len(years) >= 3 and venue_count >= 1:
            is_ref = True
            reason = f"年份数={len(years)}, 会议词={venue_count}"
        
        # 规则 6：页码模式
        elif len(page_patterns) >= 1:
            is_ref = True
            reason = f"页码模式={page_patterns}"
        
        # 规则 7：高密度 DOI/URL
        elif doi_count >= 1 or url_count >= 2:
            is_ref = True
            reason = f"DOI数={doi_count}, URL数={url_count}"
        
        # 规则 8：综合密度判断（降低阈值）
        else:
            density_score = (
                len(citations) * 2 +
                len(arxiv_refs) * 3 +
                len(author_year_matches) * 2 +
                len(years) * 0.3 +
                venue_count * 1.5 +
                len(page_patterns) * 2 +
                doi_count * 2 +
                url_count
            ) / (text_length / 100)
            
            if density_score > 1.5:  # 降低阈值
                is_ref = True
                reason = f"综合密度={density_score:.2f}"
        
        # 调试输出
        if is_ref:
            preview = text[:80].replace('\n', ' ') + '...'
            print(f"[DocumentProcessor] 🗑️ 过滤参考文献块 ({reason}): {preview}")
        
        return is_ref
    
    def split_documents(
        self,
        documents: List[Document],
        clean: bool = True
    ) -> List[Document]:
        """
        切分文档为小块
        
        Args:
            documents: 原始 Document 列表
            clean: 是否先清洗文本
            
        Returns:
            切分后的 Document 列表（chunks）
        """
        if clean:
            # 1. 【宏观截断】先尝试去掉整个参考文献章节
            documents = self.remove_references_section(documents)
            
            # 2. 清洗每个文档的文本
            for doc in documents:
                doc.page_content = self.clean_text(doc.page_content)
        
        # 使用 text_splitter 切分
        chunks = self.text_splitter.split_documents(documents)
        
        # 3. 【微观过滤】过滤掉残留的参考文献内容和垃圾块
        filtered_chunks = []
        ref_removed = 0
        garbage_removed = 0
        
        for chunk in chunks:
            content = chunk.page_content
            
            # 过滤垃圾块（Unicode 转义序列等）
            if clean and self.is_garbage_chunk(content):
                garbage_removed += 1
                continue
            
            # 过滤参考文献块
            if clean and self.is_reference_chunk(content):
                ref_removed += 1
                continue
            
            filtered_chunks.append(chunk)
        
        if ref_removed > 0 or garbage_removed > 0:
            print(f"[DocumentProcessor] ✅ 过滤完成: 参考文献块={ref_removed}, 垃圾块={garbage_removed}, 保留={len(filtered_chunks)}")
        
        # 为每个 chunk 添加索引
        for i, chunk in enumerate(filtered_chunks):
            chunk.metadata["chunk_index"] = i
        
        return filtered_chunks
    
    def process_uploaded_file(
        self,
        uploaded_file: BinaryIO,
        filename: str,
        clean: bool = True
    ) -> List[Document]:
        """
        处理上传的文件：加载 + 切分
        
        Args:
            uploaded_file: 上传的文件对象
            filename: 文件名
            clean: 是否清洗文本
            
        Returns:
            切分后的 Document 列表
        """
        # 根据文件类型选择加载方法
        if filename.lower().endswith(".pdf"):
            documents = self.load_pdf_from_upload(uploaded_file, filename)
        elif filename.lower().endswith((".docx", ".doc")):
            documents = self.load_word_from_upload(uploaded_file, filename)
        else:
            raise ValueError(f"不支持的文件类型: {filename}")
        
        # 切分文档
        chunks = self.split_documents(documents, clean=clean)
        
        return chunks